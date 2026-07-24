"""Texto de anexos Office (Excel/Word) do SEI — os que o caminho PDF/imagem descartava.

Anexo em formato Office caía em `tipo=None` no `_conteudo_via_arvore` e sumia. Planilha
de medição/faturamento e minuta Word são audit-críticas. Aqui extraímos o texto com as
libs já presentes:
  .xlsx  → openpyxl        .xls  → xlrd        .docx → python-docx
O .doc binário antigo (OLE2 Word) NÃO é suportado (precisa antiword/libreoffice, ausentes)
— devolve '' honesto, nunca explode. Detecção por magic bytes + content-type.
"""
from __future__ import annotations

import io
import logging
import shutil

logger = logging.getLogger(__name__)

_ZIP = b"PK\x03\x04"          # xlsx / docx (Office Open XML = zip)
_OLE2 = b"\xd0\xcf\x11\xe0"   # xls / doc  (formato binário antigo)
_MAX = 20000                  # teto de chars (igual ao caminho de OCR)


def _erros_parse():
    """Exceções concretas dos parsers Office (import lazy — libs pesadas)."""
    from zipfile import BadZipFile
    from openpyxl.utils.exceptions import InvalidFileException
    from xlrd.biffh import XLRDError
    from xlrd.compdoc import CompDocError
    from docx.opc.exceptions import PackageNotFoundError
    return (BadZipFile, InvalidFileException, XLRDError, CompDocError,
            PackageNotFoundError, ValueError, KeyError, OSError, TypeError, IndexError)


_ERROS_PARSE = _erros_parse()


def _do_xlsx(body: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(body), read_only=True, data_only=True)
    linhas = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cels = [str(c) for c in row if c is not None and str(c).strip()]
            if cels:
                linhas.append(" | ".join(cels))
    wb.close()
    return "\n".join(linhas)


def _do_xls(body: bytes) -> str:
    import xlrd
    wb = xlrd.open_workbook(file_contents=body)
    linhas = []
    for ws in wb.sheets():
        for r in range(ws.nrows):
            cels = [str(c.value) for c in ws.row(r) if str(c.value).strip()]
            if cels:
                linhas.append(" | ".join(cels))
    return "\n".join(linhas)


def _do_docx(body: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(body))
    partes = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tb in d.tables:
        for row in tb.rows:
            cels = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cels:
                partes.append(" | ".join(cels))
    return "\n".join(partes)


def _do_doc(body: bytes) -> str:
    """.doc binário antigo (Word 97-2003) via LibreOffice headless → txt.

    python-docx/xlrd não leem .doc antigo; o soffice converte. Perfil de usuário
    ÚNICO por chamada (permite runs concorrentes na recaptura); timeout próprio.
    """
    import subprocess
    import tempfile
    from pathlib import Path
    if not shutil.which("soffice"):
        return ""
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "in.doc"
        src.write_bytes(body)
        try:
            subprocess.run(
                ["soffice", "--headless", "--norestore",
                 f"-env:UserInstallation=file://{td}/prof",
                 "--convert-to", "txt:Text", "--outdir", td, str(src)],
                capture_output=True, timeout=120, check=False)
        except (subprocess.TimeoutExpired, OSError) as exc:
            logger.debug("office_texto: soffice falhou (%s)", str(exc)[:60])
            return ""
        out = Path(td) / "in.txt"
        return out.read_text(encoding="utf-8", errors="ignore") if out.exists() else ""


def texto_de_office(body: bytes, content_type: str = "") -> str:
    """Texto de um anexo Office; '' se não for Office suportado (nunca levanta)."""
    if not body:
        return ""
    ct = (content_type or "").lower()
    tentativas = []
    if body[:4] == _ZIP:
        # zip = xlsx OU docx. content-type desempata; senão tenta os dois.
        if "sheet" in ct or "excel" in ct:
            tentativas = [_do_xlsx, _do_docx]
        elif "word" in ct or "document" in ct:
            tentativas = [_do_docx, _do_xlsx]
        else:
            tentativas = [_do_xlsx, _do_docx]
    elif body[:4] == _OLE2:
        # ole2 = xls OU doc antigo. content-type desempata; senão tenta os dois.
        # xlrd lê .xls; .doc antigo vai pro LibreOffice (_do_doc).
        if "word" in ct or "msword" in ct:
            tentativas = [_do_doc, _do_xls]
        else:
            tentativas = [_do_xls, _do_doc]
    else:
        return ""
    for fn in tentativas:
        try:
            txt = fn(body).strip()
            if txt:
                return txt[:_MAX]
        except _ERROS_PARSE as exc:   # formato errado/corrompido/vazio → tenta o próximo
            logger.debug("office_texto: %s falhou (%s)", fn.__name__, str(exc)[:60])
    return ""
