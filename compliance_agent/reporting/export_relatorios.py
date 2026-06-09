"""
Exportador de relatórios do Hermes para .txt, .pdf e .docx.
Usa dados do banco local (OrdemBancaria, Alerta) para gerar relatórios.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

from compliance_agent.database.models import OrdemBancaria, Alerta, get_session, init_db

REPORT_DIR = Path(__file__).resolve().parent.parent / "reports"
REPORT_DIR.mkdir(exist_ok=True)


def _data_rows(limit: int | None = None):
    init_db()
    session = get_session()
    try:
        q = session.query(OrdemBancaria).order_by(
            OrdemBancaria.data_emissao.desc(), OrdemBancaria.id.desc()
        )
        if limit:
            q = q.limit(int(limit))
        return list(q)
    finally:
        session.close()


def _alert_rows(limit: int = 50):
    init_db()
    session = get_session()
    try:
        q = session.query(Alerta).order_by(Alerta.created_at.desc()).limit(int(limit))
        return list(q)
    finally:
        session.close()


def _fmt_currency(value) -> str:
    try:
        return f"R$ {float(value):,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")
    except Exception:
        return str(value)


def build_txt_report(path: Path, rows: Iterable[OrdemBancaria], alerts: Iterable[Alerta]) -> Path:
    with path.open("w", encoding="utf-8") as f:
        f.write("Relatorio de Auditoria — JFN\n")
        f.write(f"Gerado em {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(f"Total OBs: {sum(1 for _ in rows)}\n\n")
        f.write("== Alertas ==\n")
        for a in alerts:
            f.write(f"- [{a.severidade}] {a.titulo}\n")
            if a.descricao:
                f.write(f"  {a.descricao}\n")
        f.write("\n== Ordens Bancarias ==\n")
        for o in rows:
            f.write(
                f"- {o.numero_ob} | {o.favorecido_nome or 'Sem nome'} | "
                f"{_fmt_currency(o.valor)} | {o.numero_sei or 'Sem SEI'} | {o.numero_processo or '-'}\n"
            )
    return path


def build_docx_report(path: Path, rows: Iterable[OrdemBancaria], alerts: Iterable[Alerta]) -> Path:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx não está instalado. Instale com: pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading("Relatório de Auditoria — JFN", 0)
    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total OBs: {sum(1 for _ in rows)}")

    doc.add_heading("Alertas", level=1)
    for a in alerts:
        doc.add_paragraph(f"[{a.severidade}] {a.titulo}", style="List Bullet")
        if a.descricao:
            doc.add_paragraph(a.descricao)

    doc.add_heading("Ordens Bancárias", level=1)
    for o in rows:
        doc.add_paragraph(
            f"{o.numero_ob} | {o.favorecido_nome or 'Sem nome'} | "
            f"{_fmt_currency(o.valor)} | {o.numero_sei or 'Sem SEI'} | {o.numero_processo or '-'}",
            style="List Bullet",
        )

    doc.save(str(path))
    return path


def build_pdf_report(path: Path, rows: Iterable[OrdemBancaria], alerts: Iterable[Alerta]) -> Path:
    try:
        from fpdf import FPDF
    except Exception as e:
        raise RuntimeError("fpdf2 não está instalado. Instale com: pip install fpdf2") from e

    class Pdf(FPDF):
        def header(self):
            self.set_font("Arial", "B", 12)
            self.cell(0, 10, "Relatorio de Auditoria — JFN", ln=True)
            self.ln(4)

    pdf = Pdf()
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 8, f"Gerado em {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
    pdf.ln(4)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(0, 8, "Alertas", ln=True)
    pdf.set_font("Arial", "", 10)
    for a in alerts:
        txt = f"[{a.severidade}] {a.titulo} - {a.descricao or ''}"
        pdf.multi_cell(0, 8, txt)

    pdf.ln(4)
    pdf.set_font("Arial", "B", 10)
    pdf.cell(0, 8, "Ordens Bancarias", ln=True)
    pdf.set_font("Arial", "", 9)
    for o in rows:
        txt = (
            f"{o.numero_ob} | {o.favorecido_nome or 'Sem nome'} | "
            f"{_fmt_currency(o.valor)} | {o.numero_sei or 'Sem SEI'} | {o.numero_processo or '-'}"
        )
        pdf.multi_cell(0, 6, txt)

    pdf.output(str(path))
    return path


def generate_report(limit: int | None = None, fmt: str = "txt") -> dict:
    fmt = (fmt or "txt").lower().strip()
    if fmt not in {"txt", "pdf", "docx"}:
        raise ValueError("Formato inválido. Use txt, pdf ou docx.")
    rows = _data_rows(limit=limit)
    alerts = _alert_rows()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = REPORT_DIR / f"relatorio_auditoria_{stamp}.{fmt}"
    if fmt == "txt":
        build_txt_report(out, rows, alerts)
    elif fmt == "docx":
        build_docx_report(out, rows, alerts)
    elif fmt == "pdf":
        build_pdf_report(out, rows, alerts)
    return {
        "ok": True,
        "arquivo": str(out),
        "formato": fmt,
        "obs": sum(1 for _ in rows),
        "alertas": sum(1 for _ in alerts),
    }
