# -*- coding: utf-8 -*-
"""Instrumentos de mandato — JFN 2.0, Onda 10. Do achado à MINUTA parlamentar.

Fecha o ciclo: um achado/parecer vira a peça concreta — requerimento (ALERJ), representação
(TCE-RJ), notícia de fato (MP-RJ) ou post de transparência. Gera .docx (python-docx).

Invariante jurídico (cláusula de honestidade preservada): linguagem de DILIGÊNCIA/REPRESENTAÇÃO,
solicitando apuração — NUNCA afirma crime/condenação (isso compete a TCE-RJ/MP-RJ/Judiciário após
contraditório; presunção de legitimidade dos atos administrativos).
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

_OUT = Path(__file__).resolve().parent.parent / "reports"

_TIPOS = {
    "requerimento": "REQUERIMENTO DE INFORMAÇÃO (ALERJ)",
    "representacao": "REPRESENTAÇÃO AO TRIBUNAL DE CONTAS DO ESTADO DO RIO DE JANEIRO",
    "noticia_fato": "NOTÍCIA DE FATO AO MINISTÉRIO PÚBLICO DO ESTADO DO RIO DE JANEIRO",
    "post": "NOTA DE TRANSPARÊNCIA (comunicação pública)",
}

_CLAUSULA = ("Os elementos a seguir constituem INDÍCIOS que justificam DILIGÊNCIA/APURAÇÃO, à luz "
             "do dever de fiscalização do mandato parlamentar (CF/88 art. 37 e art. 70-71). NÃO se "
             "afirma a prática de ilícito: vigora a presunção de legitimidade dos atos administrativos, "
             "e a caracterização de irregularidade compete aos órgãos de controle após o contraditório.")


def _digits(s) -> str:
    return re.sub(r"\D", "", str(s or ""))


def _contexto(base: str) -> dict:
    """Resolve a base (CNPJ → resumo de OB/conflito; senão texto livre). Best-effort, honesto."""
    ctx = {"base": base, "linhas": []}
    cnpj = _digits(base)
    if len(cnpj) == 14:
        ctx["cnpj"] = cnpj
        try:
            from compliance_agent.dossie import _resumo_ob
            ob = _resumo_ob(cnpj)
            if ob.get("total_ob"):
                ctx["linhas"].append(
                    f"Pagamentos (Ordens Bancárias) ao CNPJ {cnpj}: R$ {ob['total_ob']:,.2f} em "
                    f"{ob.get('n_ob', 0)} OBs, concentração na principal UG de "
                    f"{round((ob.get('concentracao_top_ug') or 0) * 100)}%.")
        except Exception:  # noqa: BLE001
            pass
        try:
            from compliance_agent.lex_conflito import conflito
            c = conflito(cnpj=cnpj, limite=5)
            if c.get("rede"):
                ctx["linhas"].append(
                    f"Indício de conflito doador↔contrato: {len(c['rede'])} vínculo(s) entre doações "
                    f"eleitorais (TSE) e o fornecedor/sócios (a verificar).")
        except Exception:  # noqa: BLE001
            pass
    if not ctx["linhas"]:
        ctx["linhas"].append(str(base))
    return ctx


def gerar(tipo: str, base: str, precedente: bool = True) -> dict:
    """Gera a minuta .docx do instrumento. Retorna {ok, tipo, path_docx, texto}.

    tipo ∈ {requerimento, representacao, noticia_fato, post}; base = CNPJ/UG/processo/achado.
    """
    tipo = (tipo or "").strip().lower()
    if tipo not in _TIPOS:
        return {"ok": False, "erro": f"tipo inválido (use {sorted(_TIPOS)})"}
    ctx = _contexto(base)

    # corpo textual (também retornado p/ pré-visualização)
    linhas = [_TIPOS[tipo], "", _CLAUSULA, "", "DOS FATOS E INDÍCIOS:"]
    for i, ln in enumerate(ctx["linhas"], 1):
        linhas.append(f"{i}. {ln}")
    linhas += ["", "DO PEDIDO:", _pedido(tipo, ctx)]
    if precedente:
        prec = _precedente(tipo)
        if prec:
            linhas += ["", "PRECEDENTE/FUNDAMENTO:", prec]
    texto = "\n".join(linhas)

    path = _escrever_docx(tipo, texto)
    return {"ok": True, "tipo": tipo, "path_docx": path, "texto": texto,
            "_nota": "Minuta de DILIGÊNCIA/REPRESENTAÇÃO (indícios, nunca acusação). Revisar antes de protocolar."}


def _pedido(tipo: str, ctx: dict) -> str:
    alvo = ctx.get("cnpj") or ctx.get("base")
    if tipo == "requerimento":
        return (f"Requer-se ao Poder Executivo, na forma regimental, informações e documentos sobre as "
                f"contratações e pagamentos relativos a {alvo}, incluindo processos, notas de empenho, "
                f"liquidação e ordens bancárias, para o exercício do controle externo.")
    if tipo == "representacao":
        return (f"Representa-se ao TCE-RJ para que, no exercício de sua competência (CE-RJ arts. 122-123), "
                f"apure os indícios apontados quanto a {alvo}, adotando as medidas cabíveis.")
    if tipo == "noticia_fato":
        return (f"Leva-se ao conhecimento do MP-RJ a presente notícia de fato sobre {alvo}, para apuração "
                f"da eventual existência de irregularidade, sem prejuízo do contraditório.")
    return (f"Comunica-se à população, com transparência e responsabilidade, o acompanhamento das "
            f"contratações públicas relativas a {alvo}, em linguagem de diligência (sem juízo condenatório).")


def _precedente(tipo: str) -> str:
    """Cita um precedente próximo via LexML/jurisprudência (best-effort, sem inventar)."""
    try:
        from compliance_agent.lex import _run_coro
        from compliance_agent.collectors.lexml_fetcher import buscar_lexml_jurisprudencia
        res = _run_coro(lambda: buscar_lexml_jurisprudencia("sobrepreço dispensa licitação", "TCE")) or []
        if res:
            r = res[0]
            return f"{r.get('titulo', '')[:160]} ({r.get('fonte', 'LexML/TCE')})."
    except Exception:  # noqa: BLE001
        pass
    return ("Lei 14.133/2021 (arts. 5º, 9º, 18, 74-75, 125-126); Lei 8.666/93 (arts. 89-96); "
            "jurisprudência TCU/TCE-RJ sobre direcionamento e sobrepreço (consultar acórdão específico).")


def _escrever_docx(tipo: str, texto: str) -> str:
    from docx import Document

    doc = Document()
    for i, ln in enumerate(texto.split("\n")):
        if i == 0:
            doc.add_heading(ln, level=0)
        elif ln.endswith(":") and ln.isupper():
            doc.add_heading(ln, level=2)
        else:
            doc.add_paragraph(ln)
    doc.add_paragraph("")
    doc.add_paragraph(f"Minuta gerada pelo JFN em {datetime.now().strftime('%d/%m/%Y')} — revisar antes de protocolar.")
    _OUT.mkdir(parents=True, exist_ok=True)
    destino = _OUT / f"minuta_{tipo}_{datetime.now().date()}.docx"
    doc.save(str(destino))
    return str(destino)
